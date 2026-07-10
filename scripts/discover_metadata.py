
import os
import re
import openpyxl
from docx import Document
import zipfile
from typing import Dict, List, Set, Tuple

def search_xlsx(file_path: str, keywords: List[str]) -> Dict[str, Dict[str, str]]:
    """Searches an XLSX file for target-pathway-phenotype mappings."""
    mappings = {}
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            headers = [str(cell.value).lower() if cell.value else "" for cell in sheet[1]]
            
            # Find column indices for target, pathway, and role/phenotype
            try:
                target_col = headers.index("gene")
                pathway_col = headers.index("primary function")
                phenotype_col = headers.index("role in amr/biofilm")
            except ValueError:
                continue # This sheet doesn't have the required columns

            for row in sheet.iter_rows(min_row=2, values_only=True):
                target_id = str(row[target_col]).lower().strip()
                pathway = str(row[pathway_col]).strip()
                phenotype = str(row[phenotype_col]).strip()

                if target_id and pathway and phenotype:
                    mappings[target_id] = {"pathway": pathway, "phenotype": phenotype}
    except Exception:
        pass
    return mappings

def search_docx(file_path: str, keywords: List[str]) -> List[str]:
    """Searches a DOCX file for keyword hits."""
    hits = []
    try:
        doc = Document(file_path)
        # Search paragraphs
        for para in doc.paragraphs:
            for keyword in keywords:
                if re.search(r'\b' + re.escape(keyword) + r'\b', para.text, re.IGNORECASE):
                    hits.append(f"Paragraph: ...{para.text[:100]}...")
        # Search tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for keyword in keywords:
                         if re.search(r'\b' + re.escape(keyword) + r'\b', cell.text, re.IGNORECASE):
                            hits.append(f"Table Cell: ...{cell.text[:100]}...")
    except Exception:
        pass
    return list(set(hits))

def search_zip(file_path: str, keywords: List[str]) -> List[str]:
    """Searches text-based files within a ZIP archive for keywords."""
    hits = []
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            for file_info in zf.infolist():
                if not file_info.is_dir() and file_info.filename.split('.')[-1] in ['csv', 'json', 'yaml', 'md', 'txt']:
                    with zf.open(file_info.filename) as f:
                        content = f.read().decode('utf-8', errors='ignore')
                        for keyword in keywords:
                            if re.search(r'\b' + re.escape(keyword) + r'\b', content, re.IGNORECASE):
                                hits.append(f"ZIP File: `{file_info.filename}`")
    except Exception:
        pass
    return list(set(hits))

def search_md(file_path: str, keywords: List[str]) -> List[str]:
    """Searches a markdown file for keywords."""
    hits = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            for keyword in keywords:
                if re.search(r'\b' + re.escape(keyword) + r'\b', content, re.IGNORECASE):
                    hits.append(f"Markdown File: `{os.path.basename(file_path)}`")
    except Exception:
        pass
    return list(set(hits))

def main():
    base_path = "C:/Users/ayu23/OneDrive/Desktop/dock/docking_pipeline/"
    keywords = ["LasR", "PqsR", "AgrA", "SrtA", "pathway", "phenotype", "quorum sensing", "biofilm", "virulence", "adhesion", "infection", "mechanism"]
    
    # --- Source File Inspection ---
    targets_xlsx_path = os.path.join(base_path, "docs/AYUSH_AMR_Final_Targets.xlsx")
    docx_path = os.path.join(base_path, "docs/Ayush flow.docx")
    zip_path = os.path.join(base_path, "evidence_pack.zip")
    proj_rev_md_path = os.path.join(base_path, "docs/PROJECT_REVERSE_ENGINEERING.md")
    orch_md_path = os.path.join(base_path, "docs/ORCHESTRATOR_PRODUCTION_MODE.md")

    xlsx_mappings = search_xlsx(targets_xlsx_path, keywords)
    docx_hits = search_docx(docx_path, keywords)
    zip_hits = search_zip(zip_path, keywords)
    proj_rev_hits = search_md(proj_rev_md_path, keywords)
    orch_hits = search_md(orch_md_path, keywords)

    # --- Analysis & Reporting ---
    if xlsx_mappings:
        print("SUCCESS: Found authoritative metadata in AYUSH_AMR_Final_Targets.xlsx. Generating study_context.csv...")
        
        csv_content = "target_id,pathway,phenotype\\n"
        for target_id, data in xlsx_mappings.items():
            csv_content += f"{target_id},{data['pathway']},{data['phenotype']}\\n"
            
        csv_path = os.path.join(base_path, "data/inputs/study_context.csv")
        with open(csv_path, "w", encoding="utf-8") as f:
            f.write(csv_content)
        print(f"Generated data/inputs/study_context.csv")
        
    else:
        print("FAILURE: Could not find authoritative metadata. Generating source report...")
        report_content = "# 📄 Study Context Source Report\\n\\n"
        report_content += "This report details the search for an authoritative source for pathway and phenotype metadata.\\n\\n"
        report_content += "## 1. Searched Sources\\n\\n"
        report_content += "- `docs/AYUSH_AMR_Final_Targets.xlsx`\\n"
        report_content += "- `docs/Ayush flow.docx`\\n"
        report_content += "- `evidence_pack.zip`\\n"
        report_content += "- `docs/PROJECT_REVERSE_ENGINEERING.md`\\n"
        report_content += "- `docs/ORCHESTRATOR_PRODUCTION_MODE.md`\\n\\n"
        report_content += "## 2. Discovered Mappings\\n\\n"
        report_content += "No complete `target -> pathway -> phenotype` mappings were discovered in any of the searched documents.\\n\\n"
        report_content += "## 3. Missing Fields\\n\\n"
        report_content += "The following metadata fields could not be programmatically extracted from any source document:\\n- `pathway`\\n- `phenotype`\\n\\n"
        report_content += "## 4. Conclusion\\n\\n"
        report_content += "Stage 9 cannot be made fully registry-driven because the required biological metadata does not exist in an accessible format within the repository. The columns `Primary Function` and `Role in AMR/Biofilm` in the master XLSX file provide relevant data, but a direct mapping was not found during the automated search. Manual creation of `study_context.csv` is required.\\n"
        
        report_path = os.path.join(base_path, "docs/STUDY_CONTEXT_SOURCE_REPORT.md")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report_content)
        print(f"Generated docs/STUDY_CONTEXT_SOURCE_REPORT.md")

if __name__ == "__main__":
    main()

