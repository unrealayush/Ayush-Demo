
import os
import zipfile
import openpyxl
from docx import Document

def inventory_xlsx(file_path):
    """Inventories a .xlsx file."""
    inventory = {}
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            columns = [cell.value for cell in sheet[1]]
            inventory[sheet_name] = columns
    except Exception as e:
        inventory["error"] = str(e)
    return inventory

def inventory_docx(file_path):
    """Inventories a .docx file."""
    inventory = {"headings": [], "tables": []}
    try:
        document = Document(file_path)
        for para in document.paragraphs:
            if para.style.name.startswith('Heading'):
                inventory["headings"].append(para.text)
        for i, table in enumerate(document.tables):
            headers = [cell.text for cell in table.rows[0].cells]
            inventory["tables"].append({f"table_{i+1}_headers": headers})
    except Exception as e:
        inventory["error"] = str(e)
    return inventory

def inventory_zip(file_path):
    """Inventories a .zip file."""
    inventory = []
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            for file_info in zf.infolist():
                if not file_info.is_dir() and file_info.filename.split('.')[-1] in ['csv', 'json', 'yaml', 'md', 'txt']:
                    inventory.append(file_info.filename)
    except Exception as e:
        inventory.append(f"Error reading zip file: {e}")
    return inventory

def main():
    base_path = "C:/Users/ayu23/OneDrive/Desktop/dock/docking_pipeline/"
    files_to_inventory = {
        "docs/AYUSH_AMR_Final_Targets.xlsx": inventory_xlsx,
        "docs/Verified_AYUSH_Ligands_24 (1).xlsx": inventory_xlsx,
        "docs/Ayush flow.docx": inventory_docx,
        "evidence_pack.zip": inventory_zip,
    }

    report_content = "# 📦 Master Data Inventory Report

This report details the contents of the project's binary source documents.

---
"

    # Specific check for pathway/phenotype
    pathway_phenotype_source = "Not Found"

    for rel_path, inventory_func in files_to_inventory.items():
        abs_path = os.path.join(base_path, rel_path)
        report_content += f"
## 📄 File: `{rel_path}`

"
        if os.path.exists(abs_path):
            inventory = inventory_func(abs_path)
            if "error" in inventory:
                report_content += f"**Error processing file:** {inventory['error']}
"
                continue
            
            if rel_path.endswith(".xlsx"):
                for sheet, columns in inventory.items():
                    report_content += f"### Worksheet: `{sheet}`

"
                    report_content += "**Columns:**
"
                    for col in columns:
                        report_content += f"- `{col}`
"
                        if col and ('pathway' in col.lower() or 'phenotype' in col.lower()):
                            pathway_phenotype_source = f"`{rel_path}` (Sheet: `{sheet}`, Column: `{col}`)"
                    report_content += "
"
            elif rel_path.endswith(".docx"):
                report_content += "### Headings

"
                for heading in inventory["headings"]:
                    report_content += f"- {heading}
"
                report_content += "
### Tables

"
                for table in inventory["tables"]:
                    for name, headers in table.items():
                        report_content += f"- **{name}**: `{headers}`
"
            elif rel_path.endswith(".zip"):
                report_content += "### Text-Based Files Found:

"
                for fname in inventory:
                    report_content += f"- `{fname}`
"
        else:
            report_content += "**File not found.**
"
    
    report_content += "
---
"
    report_content += "## 🧬 Pathway & Phenotype Metadata

"
    report_content += f"**Authoritative Source:** {pathway_phenotype_source}

"
    
    report_content += "---
"
    report_content += """## 🔗 Dependency Graph

This graph shows the proposed authoritative source for each logical registry based on the inventory.

*   **`target_registry`**
    *   **Source:** `docs/AYUSH_AMR_Final_Targets.xlsx`
    *   **Worksheet:** `Final_Targets`
*   **`ligand_registry`**
    *   **Source:** `docs/Verified_AYUSH_Ligands_24 (1).xlsx`
    *   **Worksheet:** `Final_Ligands`
*   **`study_context`**
    *   **Source:** `docs/AYUSH_AMR_Final_Targets.xlsx` (If pathway/phenotype columns are confirmed)
    *   **Worksheet:** `Final_Targets`
*   **`assay_registry`**
    *   **Source:** Not Found
*   **`scenario_registry`**
    *   **Source:** Not Found (Logic should be derived from `target_registry` and `ligand_registry`)
"""

    report_path = os.path.join(base_path, "docs/MASTER_DATA_INVENTORY.md")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_content)
    
    print(f"Generated report at {report_path}")

if __name__ == "__main__":
    main()
